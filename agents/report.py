from pathlib import Path
from typing import Dict, Any

from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate

from rag.prompts import report_template, project_readme_prompt


def write_report(
    verdict: str,
    score: int,
    rationale: str,
    tech: str,
    market: str,
    comp: str,
    actions: str,
    sources: list[str] | None,
    candidates: list[dict] | None = None,
    path: str = "outputs/investment_report.md",
) -> str:
    cand_md = ""
    if candidates:
        lines = []
        for c in candidates:
            try:
                name = c.get("name") if isinstance(c, dict) else None  # type: ignore
                tech_txt = c.get("tech") if isinstance(c, dict) else None  # type: ignore
            except Exception:
                name, tech_txt = None, None
            if name:
                lines.append(f"- {name}: {tech_txt or ''}")
        cand_md = "\n".join(lines)

    md = report_template().format(
        verdict=verdict,
        score=score,
        rationale=rationale,
        tech=tech,
        market=market,
        comp=comp,
        candidates_tech=cand_md or "- (N/A)",
        actions=(actions or "- (none)") if verdict == "hold" else "- (N/A)",
        sources="\n".join(f"- {s}" for s in (sources or ["local index"]))
    )
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(md, encoding="utf-8")
    return str(p)


def compose_investment_brief(state: Dict[str, Any], model: str = "gpt-4o-mini") -> str:
    # Build enumerated sources and snippets
    sources = list(dict.fromkeys(state.get("sources") or []))
    src_map = {s: i + 1 for i, s in enumerate(sources)}
    snippets = state.get("snippets") or []
    snippet_block_lines = []
    for idx, sn in enumerate(snippets[:12], start=1):
        src = sn.get("src") or ""
        sid = src_map.get(src, idx)
        text = (sn.get("text") or "")
        snippet_block_lines.append(f"<<<DOC id={sid} src=\"{src}\">>>{text}<<</DOC>>")
    snippet_block = "\n".join(snippet_block_lines)

    # Context block as requested
    target = state.get("target") or ""
    domain = state.get("domain") or ""
    context_block = f"""[CONTEXT]
TARGET_COMPANY: {target}
DOMAIN: {domain}
SEGMENT: 
COUNTRY/REGION: 

PG_META: 

RAG_SNIPPETS:
{snippet_block}

COMPETITOR_CANDIDATES: {', '.join([c.get('name') for c in (state.get('candidates') or []) if isinstance(c, dict) and c.get('name')])}
ASSUMPTIONS: 최근 12~24개월 데이터 기준, 누락값은 '불명' 처리, 숫자는 단위/출처와 함께 제시
[/CONTEXT]"""

    # If we have structured market/competitor JSON, render them to Markdown first
    def _render_market(md: str, st: Dict[str, Any]) -> str:
        data = st.get("market_struct")
        if not isinstance(data, dict):
            return md
        lines = []
        ctx = data.get("context") or []
        pos = data.get("position") or []
        scores = data.get("scores") or {}
        lines.append("### 1) 산업 맥락")
        for b in (ctx if isinstance(ctx, list) else []):
            lines.append(f"- {b}")
        lines.append("\n### 2) 대상 회사 포지션")
        for b in (pos if isinstance(pos, list) else []):
            lines.append(f"- {b}")
        # scores table
        total = 0
        lines.append("\n### 3) 점수(예시)")
        for key, obj in scores.items():
            try:
                sc = int(obj.get("score", 0))
            except Exception:
                sc = 0
            rs = obj.get("reason", "")
            total += sc
            lines.append(f"- {key}: {sc}, {rs}")
        lines.append(f"**Subtotal:** {total}")
        return "\n".join(lines)

    def _render_comp(md: str, st: Dict[str, Any]) -> str:
        data = st.get("comp_struct")
        if not isinstance(data, dict):
            return md
        lines = []
        if data.get("summary"):
            lines.append(str(data["summary"]))
        headers = data.get("headers") or []
        rows = data.get("rows") or []
        if headers and rows:
            # Ensure '기준' first col exists
            hdr = [str(h) for h in headers]
            lines.append("\n### 비교표")
            lines.append("| " + " | ".join(hdr) + " |")
            lines.append("|" + "|".join(["---"] * len(hdr)) + "|")
            for r in rows:
                crit = r.get("criterion", "")
                vals = r.get("values", [])
                row = [str(crit)] + [str(v) for v in vals]
                lines.append("| " + " | ".join(row) + " |")
        if data.get("diffs"):
            lines.append("\n### 차별화 포인트(3)")
            for b in data["diffs"]:
                lines.append(f"- {b}")
        if data.get("risks"):
            lines.append("\n### 리스크/보완 과제(3)")
            for b in data["risks"]:
                lines.append(f"- {b}")
        if data.get("verdict"):
            lines.append(f"\n최종 판정: {data['verdict']}")
        return "\n".join(lines)

    market_md = _render_market(state.get("market") or "", state)
    comp_md = _render_comp(state.get("comp") or "", state)

    # Build candidates evaluation (multi-candidate results)
    cand_eval_lines: list[str] = []
    decs = state.get("decisions") or []
    if decs:
        cand_eval_lines.append("| Company | Verdict | Score |")
        cand_eval_lines.append("|---|---|---|")
        for r in decs:
            cand_eval_lines.append(f"| {r.get('name','')} | {r.get('verdict','')} | {r.get('score','')} |")
    candidates_eval = "\n".join(cand_eval_lines)

    # Build outline instruction (Korean-friendly but matching user's structure)
    outline = state.get("outline") or """# Investment Brief
## Verdict
<recommend | hold | reject> (Score <0..100>)

## Rationale
<핵심 투자 논리 3~5문장. 왜 지금? 왜 이 회사? [n] 표기 포함>

## Tech Summary
{tech}

## Market
{market}

## Competitors
{comp}

## Unit Economics & GTM
- 단건원가/그로스마진: <알려진 경우만 수치와 출처 [n]; 아니면 ‘불명’>
- CAC/LTV/회수기간: <수치 + 조건 [n] 또는 ‘불명’>
- 채널/GTM: <직판/채널/파트너/마켓플레이스 등 [n]>

## Moat & Defensibility
- 데이터/네트워크효과/프로세스 락인/규모경제 중 해당 근거 2~3개 [n]

## Team & Governance
- 핵심 경영진 약력(핵심 2~4명 요약, 직전 성과 포함) [n]
- 거버넌스/보안·컴플라이언스: <ISO, SOC2, 수출입 규정 등> [n]

## Risks & Mitigations
- <리스크 1: 영향/확률/완화전략 [n]>
- <리스크 2: 영향/확률/완화전략 [n]>
- <리스크 3: 영향/확률/완화전략 [n]>

## Investment Thesis
- 한 줄 논지: <왜 지금 이 회사인가 한 문장>
- 트리거: <다음 라운드 전 달성해야 할 정량 트리거 2~4개>

## Candidates Evaluation
{candidates_eval}

## Decision
- 최종 판정: <recommend | hold | reject> (Score <0..100>)
- 사유 요약: <불릿 3~5개 [n]>

## Next Actions (if HOLD)
- <필요 데이터/파일럿/레퍼런스/보안감사 등 구체 액션 3~5개>

## Sources
{sources_enumerated}
"""

    # Build sources enumerated
    sources_en = "\n".join([f"[{i}] {s}" for s, i in src_map.items()])

    # Assemble the full prompt to compose final report
    sys_msg = (
        "당신은 VC 투자 보고서를 작성하는 애널리스트입니다. 아래 CONTEXT와 섹션별 자료(tech/market/comp)를 활용해,"
        " 주어진 아웃라인에 딱 맞춘 Markdown만 산출하세요. 인용은 [n] 형태로 Sources의 인덱스를 참조하세요."
    )
    tmpl = ChatPromptTemplate.from_messages(
        [
            ("system", sys_msg),
            (
                "human",
                """
{context}

[SECTIONS]
- TECH: {tech}
- MARKET: {market}
- COMP: {comp}

[OUTLINE]
{outline}

[SOURCES]
{sources_enumerated}
""",
            ),
        ]
    )
    # Pre-render outline placeholders to avoid literal {tech} etc. in output
    rendered_outline = outline.format(
        tech=state.get("tech") or "",
        market=market_md,
        comp=comp_md,
        candidates_eval=candidates_eval,
        sources_enumerated=sources_en,
    )

    llm = ChatOpenAI(model=model, temperature=0)
    return (tmpl | llm).invoke(
        {
            "context": context_block,
            "tech": state.get("tech") or "",
            "market": market_md,
            "comp": comp_md,
            "outline": rendered_outline,
            "sources_enumerated": sources_en,
            "candidates_eval": candidates_eval,
        }
    ).content


def write_docx_report(
    verdict: str,
    score: int,
    rationale: str,
    tech: str,
    market: str,
    comp: str,
    actions: str,
    sources: list[str] | None,
    path: str = "outputs/investment_report.docx",
) -> str | None:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return None
    doc = Document()
    doc.add_heading("Investment Brief", level=1)
    doc.add_heading("Verdict", level=2)
    doc.add_paragraph(f"{verdict} (Score {score})")
    doc.add_heading("Rationale", level=2)
    doc.add_paragraph(rationale)
    doc.add_heading("Tech Summary", level=2)
    doc.add_paragraph(tech)
    doc.add_heading("Market", level=2)
    doc.add_paragraph(market)
    doc.add_heading("Competitors", level=2)
    doc.add_paragraph(comp)
    doc.add_heading("Next Actions", level=2)
    doc.add_paragraph((actions or "- (none)") if verdict == "hold" else "- (N/A)")
    doc.add_heading("Sources", level=2)
    for s in (sources or ["local index"]):
        doc.add_paragraph(str(s))
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(p))
    return str(p)


def generate_project_readme_md(state: Dict[str, Any], model: str = "gpt-4o-mini") -> str:
    # Deterministic project README (presentation/overview), not the investment report
    domain = state.get("domain") or ""
    query = state.get("query") or ""
    target = state.get("target") or ""
    decision = state.get("decision") or ""
    score = state.get("score") or ""
    candidates = state.get("candidates") or []
    cand_names = [c.get("name") if isinstance(c, dict) else str(c) for c in candidates]
    decs = state.get("decisions") or []
    sources = list(dict.fromkeys(state.get("sources") or []))

    # Decisions table (if available)
    dec_table = ""
    if decs:
        lines = ["| Company | Verdict | Score |", "|---|---|---|"]
        for r in decs:
            lines.append(f"| {r.get('name','')} | {r.get('verdict','')} | {r.get('score','')} |")
        dec_table = "\n".join(lines)

    tech_stack = (
        "| Category | Details |\n"
        "|---|---|\n"
        "| Framework | LangGraph, LangChain, Python |\n"
        "| LLM | GPT-4o-mini via OpenAI API |\n"
        "| Vector DB | Chroma (persisted at .index) |\n"
        "| Embedding | HuggingFace intfloat/multilingual-e5-base |\n"
        "| Database | PostgreSQL (runs/startups/sources) |\n"
        "| Tracing | LangSmith (optional) |\n"
    )

    body = f"""
# AI Startup Investment Evaluation Agent
본 프로젝트는 인공지능을 활용하여 스타트업의 투자 가능성을 자동으로 평가하는 에이전트를 설계/구현한 실습 프로젝트입니다.

## Overview
- Objective: AI 스타트업의 기술력, 시장성, 리스크 등을 기준으로 투자 적합성 분석
- Method: AI Agent + Agentic RAG + LangGraph
- Flow: Startup Search → Tech Summary → Market → Competitors → Decision → Report

## Features
- 문서 기반 RAG(Chroma) + 다국어 임베딩을 사용한 근거 중심 평가
- 후보 3개 동시 평가 및 추천 리스트 산출(Candidates Evaluation)
- Postgres에 실행/후보/출처 자동 로깅(startups, startup_sources, invest_runs)
- LangGraph 스트리밍 + tqdm 진행 표시, LangSmith 트레이싱(옵션)

## Tech Stack
{tech_stack}

## Agents
- startup_search: 도메인/쿼리 기반 후보 3개 수집(JSON)
- tech_summary: 타깃/후보 기술 요약(JSON; DB tech_raw + RAG 스니펫)
- market_eval: 산업 맥락/점수(JSON)
- competitor_analysis: 후보 3개 비교표(JSON)
- investment_decision: JSON 기반 다중 후보 의사결정
- report_writer: 보고서(investment_report.md) 및 그래프/README 생성

## Architecture
그래프 이미지는 `outputs/graph.png`에 저장됩니다. (실행 시 `--viz` 옵션)

## Directory Structure
- data/            # 문서(RAG 소스)
- agents/          # 에이전트 체인
- prompts/         # 시스템/설정 프롬프트
- rag/             # 로더/임베딩/벡터DB 유틸
- graph/app.py     # LangGraph 파이프라인
- outputs/         # 보고서/그래프 이미지

## Run Summary
- Domain: {domain}
- Query: {query}
- Candidates: {', '.join(filter(None, cand_names))}
- Decision: {decision} (Score {score})

### Candidates Evaluation
{dec_table if dec_table else '- (N/A)'}

## Sources
- 총 {len(sources)}건 참조
{os.linesep.join(['- ' + s for s in sources[:12]])}

## Contributors
- (예시) 김A: Prompt Engineering, Agent Design
- (예시) 최B: PDF Parsing, Retrieval Agent
"""
    return body


def write_text(path: str, content: str) -> str:
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(content, encoding="utf-8")
    return str(p)
